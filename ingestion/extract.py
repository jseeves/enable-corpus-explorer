"""Extract text from PDFs (and DOCX), preserving page numbers.

Output format per document:
    {
        "resource_id": "ks_001",
        "file_name": "...",
        "pages": [
            {"page_num": 1, "text": "..."},
            {"page_num": 2, "text": "..."},
            ...
        ]
    }

Page-level tracking matters because we propagate page numbers through chunking
into the vector index, so citations can point to specific pages.
"""
from __future__ import annotations

import json
from dataclasses import dataclass, field
from pathlib import Path

import fitz  # PyMuPDF


@dataclass
class ExtractedDocument:
    resource_id: str
    file_name: str
    file_path: str
    pages: list[dict] = field(default_factory=list)
    char_count: int = 0
    word_count: int = 0
    page_count: int = 0

    def to_dict(self) -> dict:
        return {
            "resource_id": self.resource_id,
            "file_name": self.file_name,
            "file_path": self.file_path,
            "pages": self.pages,
            "char_count": self.char_count,
            "word_count": self.word_count,
            "page_count": self.page_count,
        }


def extract_pdf(path: Path, resource_id: str) -> ExtractedDocument:
    """Extract text from a PDF, page by page."""
    doc = fitz.open(path)
    pages = []
    for i in range(doc.page_count):
        text = doc[i].get_text()
        # Light cleanup: collapse runs of whitespace, strip page-edge noise
        text = "\n".join(line.strip() for line in text.splitlines() if line.strip())
        pages.append({"page_num": i + 1, "text": text})
    doc.close()

    full_text = "\n".join(p["text"] for p in pages)
    return ExtractedDocument(
        resource_id=resource_id,
        file_name=path.name,
        file_path=str(path),
        pages=pages,
        char_count=len(full_text),
        word_count=len(full_text.split()),
        page_count=len(pages),
    )


def extract_docx(path: Path, resource_id: str) -> ExtractedDocument:
    """Extract text from a DOCX. Pages are not preserved (DOCX has no fixed pagination);
    we treat the whole doc as a single 'page' for compatibility."""
    from docx import Document  # local import; only needed for docx files

    d = Document(path)
    text = "\n".join(p.text.strip() for p in d.paragraphs if p.text.strip())
    return ExtractedDocument(
        resource_id=resource_id,
        file_name=path.name,
        file_path=str(path),
        pages=[{"page_num": 1, "text": text}],
        char_count=len(text),
        word_count=len(text.split()),
        page_count=1,
    )


def extract_one(path: Path, resource_id: str) -> ExtractedDocument:
    """Dispatch on file extension."""
    suffix = path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf(path, resource_id)
    elif suffix == ".docx":
        return extract_docx(path, resource_id)
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def write_extracted(docs: list[ExtractedDocument], out_path: Path) -> None:
    """Write all extracted docs as JSONL for downstream chunking."""
    out_path.parent.mkdir(parents=True, exist_ok=True)
    with out_path.open("w") as f:
        for d in docs:
            f.write(json.dumps(d.to_dict()) + "\n")


if __name__ == "__main__":
    # Quick smoke test
    import sys
    from pathlib import Path

    if len(sys.argv) < 2:
        print("Usage: python -m ingestion.extract <path-to-pdf>")
        sys.exit(1)

    p = Path(sys.argv[1])
    doc = extract_one(p, "test_001")
    print(f"Extracted {doc.page_count} pages, {doc.word_count} words from {doc.file_name}")
    print(f"\nFirst 500 chars:\n{doc.pages[0]['text'][:500]}")
